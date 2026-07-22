import os
import sys
import webbrowser
import subprocess
import urllib.parse

sys.dont_write_bytecode = True


def desplegar_monitores_windows() -> bool:
    """Despliega la configuración de monitores nativos en Windows."""
    try:
        subprocess.Popen("displayswitch.exe /extend", shell=True)
        print("[SystemCommands]: Monitores de Windows configurados correctamente.")
        return True
    except Exception as e:
        print(f"[SystemCommands]: Error al desplegar monitores: {e}")
        return False


def buscar_en_navegador_sistema(consulta: str) -> str:
    """Abre el navegador predeterminado y busca en Google."""
    if not consulta:
        return "No se especificó ninguna consulta."
    try:
        query = urllib.parse.quote(consulta.strip())
        webbrowser.open(f"https://www.google.com/search?q={query}")
        return f"Buscando '{consulta}' en el navegador."
    except Exception as e:
        return f"Error abriendo navegador: {e}"


def reproducir_video_brave(busqueda: str) -> str:
    """Abre YouTube con la búsqueda solicitada."""
    if not busqueda:
        return "No se especificó búsqueda de video."
    try:
        query = urllib.parse.quote(busqueda.strip())
        webbrowser.open(f"https://www.youtube.com/results?search_query={query}")
        return f"Buscando video '{busqueda}' en YouTube."
    except Exception as e:
        return f"Error al abrir YouTube: {e}"


def crear_y_abrir_documento_word(nombre_archivo: str, tema_o_contenido: str, carpeta_destino: str = None) -> str:
    """Crea un archivo .docx REAL con contenido sobre el tema solicitado y lo abre en Microsoft Word."""
    try:
        if carpeta_destino and os.path.isabs(carpeta_destino):
            ruta_dir = carpeta_destino
        elif carpeta_destino:
            ruta_dir = os.path.join(os.path.expanduser("~"), "Desktop", carpeta_destino)
        else:
            ruta_dir = os.path.join(os.path.expanduser("~"), "Desktop")

        os.makedirs(ruta_dir, exist_ok=True)

        if not nombre_archivo.endswith(".docx"):
            nombre_archivo += ".docx"

        ruta_completa = os.path.join(ruta_dir, nombre_archivo)

        try:
            import docx
            doc = docx.Document()
            doc.add_heading(nombre_archivo.replace(".docx", ""), level=1)
            doc.add_paragraph(f"Documento generado por REVAN sobre la temática: {tema_o_contenido}\n")
            doc.add_paragraph(tema_o_contenido)
            doc.save(ruta_completa)
        except ImportError:
            with open(ruta_completa.replace(".docx", ".txt"), "w", encoding="utf-8") as f:
                f.write(f"--- {nombre_archivo} ---\n\nTema: {tema_o_contenido}")
            ruta_completa = ruta_completa.replace(".docx", ".txt")

        os.startfile(ruta_completa)
        return f"Archivo '{nombre_archivo}' creado con éxito con la información sobre '{tema_o_contenido}' y abierto en pantalla."

    except Exception as e:
        return f"Error al crear el documento: {e}"


def lanzar_aplicacion_usuario(nombre_app: str) -> str:
    """Lanza aplicaciones populares o ejecutables de Windows."""
    nombre = nombre_app.lower().strip()
    try:
        if any(k in nombre for k in ["calc", "calculadora"]):
            subprocess.Popen("calc.exe")
            return "Calculadora abierta."
        elif any(k in nombre for k in ["bloc", "notepad", "notas"]):
            subprocess.Popen("notepad.exe")
            return "Bloc de notas abierto."
        elif any(k in nombre for k in ["brave", "chrome", "navegador", "internet"]):
            webbrowser.open("https://www.google.com")
            return "Navegador abierto."
        else:
            os.system(f"start {nombre_app}")
            return f"Ejecutando {nombre_app}."
    except Exception as e:
        return f"Error al lanzar la aplicación {nombre_app}: {e}"


def lanzar_videojuego(nombre_juego: str) -> str:
    """Lanza el juego de Minecraft u otros usando el protocolo nativo de Windows o el launcher."""
    nombre = nombre_juego.lower().strip()
    try:
        if "minecraft" in nombre:
            os.system("start minecraft:")
            return "Iniciando Minecraft."
        else:
            os.system(f"start {nombre_juego}")
            return f"Iniciando {nombre_juego}."
    except Exception as e:
        return f"Error al intentar abrir el juego: {e}"


def ejecutar_aplicacion_office(app: str) -> str:
    """Abre aplicaciones de la suite Microsoft Office."""
    nombre = app.lower().strip()
    try:
        if "word" in nombre:
            subprocess.Popen("winword.exe")
            return "Microsoft Word iniciado."
        elif "excel" in nombre:
            subprocess.Popen("excel.exe")
            return "Microsoft Excel iniciado."
        elif "powerpoint" in nombre or "ppt" in nombre:
            subprocess.Popen("powerpnt.exe")
            return "Microsoft PowerPoint iniciado."
        else:
            os.system(f"start {app}")
            return f"Ejecutando {app}."
    except Exception as e:
        return f"Error al abrir la aplicación de Office '{app}': {e}"

__all__ = [
    "buscar_en_navegador_sistema",
    "reproducir_video_brave",
    "lanzar_aplicacion_usuario",
    "lanzar_videojuego",
    "desplegar_monitores_windows",
    "ejecutar_aplicacion_office",
    "crear_y_abrir_documento_word",
]